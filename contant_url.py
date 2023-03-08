# coding:utf-8
import re
import requests


def get_data(id, title):
    """获取数据所在页面"""
    # 检查id
    if '.html' in str(id):
        id = str(id).replace('.html', '')
    url = 'https://wenku.baidu.com/ndocview/readerinfo'

    cookies = {
        'BIDUPSID': 'C7D9963B67BA811E923F2D407E629C10',
        'PSTM': '1617448568',
        '_click_param_reader_query_ab': '-1',
        '__yjs_duid': '1_d31de959fb659416cd56b3df01ca74c71620300019707',
        'BAIDUID': '8099F3C8D6ED2F15F4829FE2620342B5:FG=1',
        'BAIDUID_BFESS': '8099F3C8D6ED2F15F4829FE2620342B5:FG=1',
        'ZFY': 'UNhchDEPpAGJHDUVC9lcP:AyTGZ2:Ae9:AETENHvxkO7mw:C',
        '__bid_n': '1842c6732ebb26fd8f4207',
        'FPTOKEN': 'oeuYR2uWCFoQcVr7ryaH/JzqYzuE2gX+upC2kTv80KedKGSeC2UVExknlBFC+AsQe5TlzLzna5S/KJq41b75E2jMhV+5DHBw+wUsz8g8LEiFmSqDJS9HOJ7Pl3eV8IbVFvbGUNGK11Zhen2FcZoMDzPutxtzZ6q2c/XyOkw9DFWcLywFX5NIGxZG3mjlSwmvtYAZJyK5EREHasuYSfog890IxDTfX2KC/6L7yQHXbH+Y9UXD4K2EQqDtSgkgfqCOZQR5bdfEhscEtl3Zh2k8je1acRX2unvW6kjAk4BNcWs1VCBpA/Y6GqiRcGuVlEH/cpBReVhB2iu4RD5BQ+kCtvMxUgVZGaHMaJ5KG+D1xsLM2cwsEHMK/pfaQovHKfg+40PwLA9/+EvZpgdew5j6gA==|Ue9SVTxwfKKsTI7RpUSGZ/zVmJkq25iHE/jFuQ/cRVA=|10|3e4bb7ad8887e5b62ffa25c7eadc1ed9',
        'BA_HECTOR': '0l2ka5818h80ak2hag0kal901i09fhi1m',
        'BDORZ': 'B490B5EBF6F3CD402E515D22BCDA1598',
        '___wk_scode_token': 'JUGDnxyCONDIzvMvGu7hqzl9Fi59GqEYMrHRo2xwTdg%3D',
        'view_edit_mode_guide': '1',
        'Hm_lvt_d8bfb560f8d03bbefc9bdecafc4a4bf6': '1676109941,1676285047,1678027757,1678090078',
        'Hm_lvt_f06186a102b11eb5f7bcfeeab8d86b34': '1676286180,1678027816,1678092322',
        'LoseUserAllPage': '%7B%22status%22%3A0%2C%22expire_time%22%3A0%2C%22create_time%22%3A1678092492%2C%22type%22%3A0%2C%22cookie_time%22%3A1678351692%7D',
        'Hm_lpvt_f06186a102b11eb5f7bcfeeab8d86b34': '1678092493',
        'Hm_lpvt_d8bfb560f8d03bbefc9bdecafc4a4bf6': '1678092567',
        'ab_sr': '1.0.1_Yzk0ODVjMzRkODExOTdjODM2NDBlOWE2MjJlZmNkZjViYzA0N2ZhZDdlOTY4YmFmOTk4ZDczMjJhYTc2NDUxZWQxYjQwYjVkYzZhNTQ2NDEyYjk0YzQ0MThjMTkwNzY3ZjQwZTA3MWJjZjc3OWEwYmRjZTQ3YzY3Mjc0NDE1ODFmNzI5ZmFhNGVmYzZiMDhkZDkxMGFjN2UzMTU2NmJiNw==',
        'bcat': '9ad7079b1d315f7c06b47714df9bcbe5166d91253da019990ce6c5f182df7431572306e72c957580cbebad2e73df2c5566338d7aef9d59d018409fe30cc9cc959ca59bc5452abdad2ec51ab1f7f7dbf4256a3c437437bfb3b8b3aa847ffe0204',
    }

    headers = {
        'Accept': 'application/json, text/plain, */*',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'Cache-Control': 'no-cache',
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Referer': 'https://wenku.baidu.com/view/01bd86e383c4bb4cf7ecd1e1.html?fr=hp_doclist&_wkts_=1678027950384',
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'User-Agent': "Mozilla / 5.0(Windows NT 10.0; Win64; x64) AppleWebKit / 537.36(KHTML, like Gecko) Chrome / 80.0.3987.122  Safari / 537.36",
        'sec-ch-ua': '\\',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '\\',
    }
    params = {
        'doc_id': id,
        'docId': id,
        'type': 'html',
        'clientType': '100',  # 这个值改为100表示可以获取100页的数据
        'pn': '0',  # 这个只有为0可以获取全部数据
        't': '1678177299733',
        'isFromBdSearch': '0',
        'srcRef': '',
        'rn': '50',
        'powerId': '2',
        'bizName': 'mainPc',
    }

    response = requests.get(url, headers=headers, params=params, cookies=cookies)
    url = []
    # 获取可以找到数据的网址
    for i in response.json()['data']['htmlUrls']['json']:
        url.append(str(i['pageLoadUrl']).replace('\\', ''))
    # # 获得页面数据
    for i in url:
        data = requests.get(i, headers=headers)
        data = ("Response:", data.text.encode('utf-8').decode('unicode_escape'))
        print(data)
        write_word(body_parse(data), style_parse(data), title)


def style_parse(data):
    """样式解析"""
    data = re.findall('"style":\[(.*?)\],"body"', str(data))[0]
    data = re.findall('{"t":(.*?),"c":(.*?),"s":{(.*?)}', data)  # 获取数据的样式
    style = {}
    # 设置每个编号的格式
    for tyle, id, type in data:
        for i in str(id).replace('[', '').replace(']', '').split(','):  # 分割字体id
            ##创建存在编号的集合
            if i in style:
                pass
            else:
                style[i] = {}

            for first_j in str(type).split(','):
                key = str(first_j).split(':')[0]
                value = str(first_j).split(':')[1]
                style[i][key] = value
    return style


def body_parse(data):
    '''内容解析'''
    data = re.findall('"body":\[(.*?)\],"page"', str(data))[0]
    data = re.findall('{"c":"(.*?)","p":{(.*?)},"ps":.*?,"t":"(.*?)","r":\[(.*?)\]}', data)
    body_data = []
    for contant, pos, type, style in data:
        pos = re.findall('"y":(.*?),', pos)[0]
        body_data.append([contant, pos, style])
    return body_data


data = ('Response:',
        'wenku_1({"outline":null,"outlineMiss":null,"font":{"1822fee6f90f76c660371a090010001":"TimesNewRomanPSMT","1822fee6f90f76c660371a090020001":"SimHei","1822fee6f90f76c660371a090030001":"FangSong"},"style":[{"t":"style","c":[1,2,3,5,6,7,8,9,10,11,12,13,0],"s":{"color":"#000000"}},{"t":"style","c":[1],"s":{"font-family":"1822fee6f90f76c660371a090010001","font-size":"13.5"}},{"t":"style","c":[2],"s":{"font-family":"1822fee6f90f76c660371a090020001","font-size":"27"}},{"t":"style","c":[2,6,3],"s":{"bold":"true"}},{"t":"style","c":[5,6,7,8,9,10,11,12,13,4],"s":{"font-family":"1822fee6f90f76c660371a090030001","font-size":"21.06"}},{"t":"style","c":[6,7,8,9,10,11,12,13,5],"s":{"color":"#000000"}},{"t":"style","c":[6],"s":{"bold":"true"}},{"t":"style","c":[7],"s":{"letter-spacing":"0.09"}},{"t":"style","c":[8],"s":{"letter-spacing":"-0.063"}},{"t":"style","c":[9],"s":{"letter-spacing":"-0.065"}},{"t":"style","c":[10],"s":{"letter-spacing":"-0.051"}},{"t":"style","c":[11],"s":{"letter-spacing":"-0.089"}},{"t":"style","c":[12],"s":{"letter-spacing":"-0.056"}},{"t":"style","c":[13],"s":{"letter-spacing":"-0.081"}}],"body":[{"c":" ","p":{"h":14.525,"w":3.375,"x":446.475,"y":64.903,"z":0},"ps":{"_enter":1},"t":"word","r":[1]},{"c":"    ","p":{"h":26.999,"w":54.161,"x":252.93,"y":117.956,"z":3},"ps":null,"s":{"letter-spacing":"0.053"},"t":"word","r":[2]},{"c":"六级数学上册教材分析","p":{"h":26.999,"w":271.044,"x":307.335,"y":117.956,"z":4},"ps":null,"s":{"letter-spacing":"0.115"},"t":"word","r":[2]},{"c":{"ix":0,"iy":0,"iw":630,"ih":39},"p":{"h":39,"w":630,"x":132,"y":185,"z":5},"ps":{"_drop":1,"_vector":1},"s":null,"t":"pic"},{"c":"  ","p":{"h":26.999,"w":27,"x":578.444,"y":117.956,"z":6},"ps":{"_enter":1},"t":"word","r":[2]},{"c":"一、学情分析","p":{"h":21.06,"w":126.528,"x":135.036,"y":188.759,"z":7},"ps":null,"s":{"letter-spacing":"0.033"},"t":"word","r":[6]},{"c":" ","p":{"h":21.06,"w":10.529,"x":261.57,"y":188.759,"z":8},"ps":{"_enter":1},"s":{"color":"#656565"},"t":"word","r":[4]},{"c":"1.","p":{"h":21.06,"w":21.15,"x":176.97,"y":251.939,"z":9},"ps":null,"t":"word","r":[7]},{"c":"本班男女生人数基本均衡，学生总体比较调皮，活泼，上课气","p":{"h":21.06,"w":559.901,"x":198.03,"y":251.939,"z":10},"ps":null,"s":{"letter-spacing":"-0.335"},"t":"word","r":[5]},{"c":"氛活跃。","p":{"h":21.06,"w":84.239,"x":135.036,"y":285.059,"z":11},"ps":null,"t":"word","r":[5]},{"c":"其中男生在课堂上的思维敏捷性优于女生，","p":{"h":21.06,"w":399.002,"x":215.674,"y":285.059,"z":12},"ps":null,"t":"word","r":[8]},{"c":"但学习上缺少耐","p":{"h":21.06,"w":147.04,"x":611.076,"y":285.059,"z":13},"ps":null,"t":"word","r":[8]},{"c":"心与细心，","p":{"h":21.06,"w":105.299,"x":135.036,"y":317.999,"z":14},"ps":null,"t":"word","r":[5]},{"c":"女生相对男生来说学习更加认真，","p":{"h":21.06,"w":315.141,"x":236.566,"y":317.999,"z":15},"ps":null,"s":{"letter-spacing":"-0.054"},"t":"word","r":[5]},{"c":"但分析判断能力却不及","p":{"h":21.06,"w":210.031,"x":548.106,"y":317.999,"z":16},"ps":null,"t":"word","r":[8]},{"c":"男生。","p":{"h":21.06,"w":63.179,"x":135.036,"y":350.939,"z":17},"ps":null,"t":"word","r":[5]},{"c":" ","p":{"h":21.06,"w":10.53,"x":198.03,"y":350.939,"z":18},"ps":{"_enter":1},"t":"word","r":[5]},{"c":"2.","p":{"h":21.06,"w":21.15,"x":176.97,"y":384.104,"z":19},"ps":null,"t":"word","r":[7]},{"c":"学生以往成绩：","p":{"h":21.06,"w":147.23,"x":198.03,"y":384.104,"z":20},"ps":null,"s":{"letter-spacing":"-0.031"},"t":"word","r":[5]},{"c":"学生的基础参差不齐，","p":{"h":21.06,"w":210.031,"x":341.49,"y":384.104,"z":21},"ps":null,"t":"word","r":[8]},{"c":"严重影响了学生对新课","p":{"h":21.06,"w":210.22,"x":547.92,"y":384.104,"z":22},"ps":null,"s":{"letter-spacing":"-0.042"},"t":"word","r":[5]},{"c":"程的学习，尤其是计算能力较差，部分学生连小数的乘除法都不会。","p":{"h":21.06,"w":629.904,"x":135.036,"y":417.044,"z":23},"ps":null,"t":"word","r":[9]},{"c":"据了解，学生以前成绩不是太好，尤其是上学期期末考试。另外也存","p":{"h":21.06,"w":622.891,"x":135.036,"y":449.984,"z":24},"ps":null,"s":{"letter-spacing":"-0.307"},"t":"word","r":[5]},{"c":"在着两级分化现象的现象。","p":{"h":21.06,"w":252.151,"x":135.036,"y":483.104,"z":25},"ps":null,"t":"word","r":[10]},{"c":" ","p":{"h":21.06,"w":10.529,"x":387.075,"y":483.104,"z":26},"ps":{"_enter":1},"t":"word","r":[5]},{"c":"3.","p":{"h":21.06,"w":21.15,"x":135.036,"y":516.044,"z":27},"ps":null,"t":"word","r":[7]},{"c":"学习习惯：约一半学生能够主动学习，比较喜欢上数学课，学习热","p":{"h":21.06,"w":601.831,"x":156.09,"y":516.044,"z":28},"ps":null,"s":{"letter-spacing":"-0.318"},"t":"word","r":[5]},{"c":"情也很高。","p":{"h":21.06,"w":105.299,"x":135.036,"y":548.984,"z":29},"ps":null,"t":"word","r":[5]},{"c":"约有四分之一学生学习需要经常由老师督促。","p":{"h":21.06,"w":419.856,"x":236.566,"y":548.984,"z":30},"ps":null,"s":{"letter-spacing":"-0.07"},"t":"word","r":[5]},{"c":"剩余四分之","p":{"h":21.06,"w":105.11,"x":652.821,"y":548.984,"z":31},"ps":null,"s":{"letter-spacing":"-0.047"},"t":"word","r":[5]},{"c":"一学生学习较懒散、习惯差，如粗心大意、书写不认真，不愿思考问","p":{"h":21.06,"w":623.123,"x":135.036,"y":582.104,"z":32},"ps":null,"s":{"letter-spacing":"-0.299"},"t":"word","r":[5]},{"c":"题，上课开小差，抄袭作业等。做作业时不细心，导致作业中经常出","p":{"h":21.06,"w":623.312,"x":135.036,"y":615.044,"z":33},"ps":null,"s":{"letter-spacing":"-0.292"},"t":"word","r":[5]},{"c":"现抄错数字或简单计算出错的情况，","p":{"h":21.06,"w":336.012,"x":135.036,"y":647.984,"z":34},"ps":null,"t":"word","r":[8]},{"c":"没有认真独立完成作业的学生占","p":{"h":21.06,"w":293.892,"x":464.035,"y":647.984,"z":35},"ps":null,"s":{"letter-spacing":"-0.072"},"t":"word","r":[5]},{"c":"有相当比例。","p":{"h":21.06,"w":126.17,"x":135.036,"y":681.134,"z":36},"ps":null,"s":{"letter-spacing":"-0.037"},"t":"word","r":[5]},{"c":" ","p":{"h":21.06,"w":10.529,"x":261.03,"y":681.134,"z":37},"ps":null,"t":"word","r":[5]},{"c":"二、本册教材分析","p":{"h":21.06,"w":168.642,"x":135.036,"y":718.574,"z":38},"ps":null,"s":{"letter-spacing":"0.023"},"t":"word","r":[6]},{"c":"  ","p":{"h":21.06,"w":20.97,"x":303.914,"y":718.574,"z":39},"ps":{"_enter":1},"s":{"bold":"true"},"t":"word","r":[6,11]},{"c":"    ","p":{"h":21.06,"w":42.014,"x":135.036,"y":760.874,"z":40},"ps":null,"s":{"letter-spacing":"-0.035"},"t":"word","r":[5]},{"c":"本册教科书的主要内容有：","p":{"h":21.06,"w":252.151,"x":176.97,"y":760.874,"z":41},"ps":null,"t":"word","r":[10]},{"c":" ","p":{"h":21.06,"w":10.529,"x":426.855,"y":760.874,"z":42},"ps":null,"t":"word","r":[5]},{"c":"分数乘法，位置与方向，分数除法，","p":{"h":21.06,"w":331.445,"x":437.294,"y":760.874,"z":43},"ps":null,"s":{"letter-spacing":"-0.367"},"t":"word","r":[5]},{"c":"比，圆，百分数，扇形统计图、数学广角——数与形。分数乘法和除","p":{"h":21.06,"w":623.397,"x":135.036,"y":793.814,"z":44},"ps":null,"s":{"letter-spacing":"-0.289"},"t":"word","r":[5]},{"c":"法，圆，比和按比例分配是本册教材的重点教学内容。","p":{"h":21.06,"w":503.923,"x":135.036,"y":826.934,"z":45},"ps":null,"t":"word","r":[9]},{"c":" ","p":{"h":21.06,"w":10.529,"x":638.925,"y":826.934,"z":46},"ps":{"_enter":1},"t":"word","r":[5]},{"c":"    1","p":{"h":21.06,"w":52.643,"x":135.036,"y":859.874,"z":47},"ps":null,"t":"word","r":[5]},{"c":"、在数与代数方面，教材安排了分数乘法、分数除法、分数混","p":{"h":21.06,"w":566.914,"x":187.59,"y":859.874,"z":48},"ps":null,"t":"word","r":[9]},{"c":"合运算比和按比例分配、","p":{"h":21.06,"w":231.094,"x":135.036,"y":892.814,"z":49},"ps":null,"t":"word","r":[12]},{"c":"负数的初步认识五个单元。","p":{"h":21.06,"w":252.151,"x":362.529,"y":892.814,"z":50},"ps":null,"t":"word","r":[10]},{"c":"分数乘法和除法","p":{"h":21.06,"w":147.04,"x":611.079,"y":892.814,"z":51},"ps":null,"t":"word","r":[8]},{"c":"的教学是在前面学习整数、","p":{"h":21.06,"w":252.151,"x":135.036,"y":925.979,"z":52},"ps":null,"t":"word","r":[10]},{"c":"小数有关计算的基础上，","p":{"h":21.06,"w":231.091,"x":383.586,"y":925.979,"z":53},"ps":null,"t":"word","r":[12]},{"c":"培养学生分数四","p":{"h":21.06,"w":147.04,"x":611.076,"y":925.979,"z":54},"ps":null,"t":"word","r":[8]},{"c":"则运算能力以及解决有关分数的实际问题的能力。","p":{"h":21.06,"w":461.993,"x":135.036,"y":958.919,"z":55},"ps":null,"t":"word","r":[8]},{"c":"会解决简单的有关","p":{"h":21.06,"w":167.911,"x":590.016,"y":958.919,"z":56},"ps":null,"t":"word","r":[13]},{"c":"比和按比例分配的实际问题，是小学生应具备的基本数学能力。","p":{"h":21.06,"w":587.974,"x":135.036,"y":991.859,"z":57},"ps":null,"t":"word","r":[8]},{"c":"    ","p":{"h":21.06,"w":41.849,"x":723.21,"y":991.859,"z":58},"ps":null,"t":"word","r":[11]},{"c":"2","p":{"h":21.059,"w":10.53,"x":135.036,"y":1024.979,"z":59},"ps":null,"t":"word","r":[5]},{"c":"、在空间与图形方面，教材安排了图形的圆、变换和确定位置","p":{"h":21.059,"w":560.09,"x":145.476,"y":1024.979,"z":60},"ps":null,"s":{"letter-spacing":"-0.328"},"t":"word","r":[5]},{"c":" ","p":{"h":21.059,"w":10.529,"x":705.75,"y":1024.979,"z":61},"ps":null,"t":"word","r":[5]},{"c":"两个","p":{"h":21.059,"w":42.12,"x":716.189,"y":1024.979,"z":62},"ps":null,"t":"word","r":[5]},{"c":"单元。","p":{"h":21.059,"w":63.179,"x":135.036,"y":1057.919,"z":63},"ps":null,"t":"word","r":[5]},{"c":"通过丰富的现实的数学活动，","p":{"h":21.059,"w":273.021,"x":189.391,"y":1057.919,"z":64},"ps":null,"t":"word","r":[8]},{"c":"让学生经历初步的数学化的过程，","p":{"h":21.059,"w":314.762,"x":453.589,"y":1057.919,"z":65},"ps":null,"t":"word","r":[13]},{"c":"理解并学会用数对表示位置；初步认识研究曲线图形的基本基本方","p":{"h":21.059,"w":608.844,"x":135.036,"y":1090.859,"z":66},"ps":null,"s":{"letter-spacing":"-0.067"},"t":"word","r":[5]},{"c":"法，促进学生空间观念的进一步发展。","p":{"h":21.059,"w":356.882,"x":135.036,"y":1123.973,"z":67},"ps":null,"s":{"letter-spacing":"-0.071"},"t":"word","r":[5]},{"c":" ","p":{"h":21.059,"w":10.529,"x":492.044,"y":1123.973,"z":68},"ps":{"_enter":1},"t":"word","r":[5]},{"c":" ","p":{"h":14.526,"w":3.375,"x":135.036,"y":1174.027,"z":1},"ps":{"_enter":1},"t":"word","r":[1]},{"c":"1 ","p":{"h":14.526,"w":10.214,"x":751.29,"y":1174.207,"z":2},"ps":{"_enter":1},"s":{"letter-spacing":"0.089"},"t":"word","r":[1]}],"page":{"ph":1262.879,"pw":892.979,"iw":630,"ih":39,"v":6,"t":"1","pptlike":false,"cx":132,"cy":64.903,"cw":636.739,"ch":1123.83,"flag":"cloud_conv_time 2022-04-24 18:11:54","ox":117,"oy":55.903,"ow":716.639,"oh":1141.83}})')


def write_word(body, style, title):
    import docx
    from docx.shared import Pt, RGBColor
    # 创建文档
    try:
        doc = docx.Document(rf'{title}.docx')
    except:
        doc = docx.Document()
    juage_y = '0'
    for contant, y, id in body:
        if juage_y != y:
            juage_y = y
            p = doc.add_paragraph('')
        t = p.add_run(contant)
        for i in str(id).split(','):
            if '"bold"' in style[i]:
                t.bold = True
    doc.save(f'{title}.docx')


def get_url(url):
    headers = {
        'user-agent': 'Mozilla / 5.0(Windows NT 10.0; Win64; x64) AppleWebKit / 537.36(KHTML, like Gecko) Chrome / 80.0.3987.122  Safari / 537.36'

    }
    title = requests.get(url, headers=headers)
    id = re.findall('view/(.*?)\?', url)
    title = re.findall('<title>(.*?)</title>', title.text)
    get_data(id[0], title[0])
    # print(title)


url = input('请输入爬取的百度文库网址：')
try:
    get_url(url)
except KeyError as e:
    print(f'失败原因：{e}')
